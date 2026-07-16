from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "downloads"
OUT.mkdir(exist_ok=True)
DOCX = OUT / "HarmonyLink_Partner_Policy_v1.0.docx"

NAVY = "082C68"
BLUE = "1155D9"
PALE = "EAF3FF"
LIGHT = "F4F7FB"
INK = "10213B"
MUTED = "65738A"
GOLD = "F3B442"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(.78)
sec.bottom_margin = Inches(.75)
sec.left_margin = Inches(.82)
sec.right_margin = Inches(.82)
sec.header_distance = Inches(.35)
sec.footer_distance = Inches(.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Malgun Gothic"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(10.3)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for style_name, size, color, before, after in [
    ("Title", 27, NAVY, 0, 8),
    ("Heading 1", 18, NAVY, 18, 8),
    ("Heading 2", 14, BLUE, 13, 6),
    ("Heading 3", 11.5, NAVY, 9, 4),
]:
    s = styles[style_name]
    s.font.name = "Malgun Gothic"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = style_name != "Title"
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_cell_text(cell, text, bold=False, color=INK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = "Malgun Gothic"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

def add_callout(label, text, fill=PALE, color=NAVY):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.7)
    c = t.cell(0,0); shade(c, fill); cell_margins(c, 150, 170, 150, 170)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(color)
    p2 = c.add_paragraph(); p2.paragraph_format.space_after=Pt(0); p2.paragraph_format.line_spacing=1.18
    r2=p2.add_run(text); r2.font.size=Pt(9.5); r2.font.color.rgb=RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)

header = sec.header
hp = header.paragraphs[0]
hp.text = "HARMONY LINK  |  PARTNER CENTER"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for r in hp.runs:
    r.font.name="Arial"; r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
fp = sec.footer.paragraphs[0]
fp.add_run("Hibelle Consulting Inc.  ·  Partner Policy v1.0")
fp.runs[0].font.size=Pt(8); fp.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
add_page_number(fp)

# Cover / customer-pack opening
p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(28); p.paragraph_format.space_after=Pt(4)
r=p.add_run("PARTNER OPERATIONS POLICY"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(BLUE); r.font.name="Arial"
p = doc.add_paragraph(style="Title"); p.add_run("HarmonyLink 입점 파트너\n플랫폼 이용 및 운영 정책")
p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(20)
r=p.add_run("Version 1.0  |  운영사 Hibelle Consulting Inc."); r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(MUTED)
add_callout("문서 목적", "입점 파트너가 HarmonyLink의 매칭 구조, 멤버십, 플랫폼 이용수수료, 수업 진행 절차와 운영 원칙을 명확히 이해하도록 안내하는 공식 운영 정책입니다.")

doc.add_heading("HarmonyLink 소개", level=1)
doc.add_paragraph("HarmonyLink(이음문화센터)는 전문 강사·교육업체와 교육기관, 단체, 소그룹 및 개인 수강자를 연결하는 방문형 문화·교육 플랫폼입니다. 운영사는 Hibelle Consulting Inc.입니다.")
add_callout("중요", "HarmonyLink는 강사를 고용하거나 직접 파견하는 교육기관이 아닙니다. 입점 파트너는 독립적으로 교육서비스를 제공하며, HarmonyLink는 검토된 파트너의 프로그램 정보 제공, 매칭, 일정 조율과 운영 지원을 담당합니다.", fill="FFF5DD", color="8A5A00")

doc.add_heading("목차", level=1)
toc = [
    "1. 운영 목적", "2. 플랫폼 운영 방식", "3. 입점 파트너 자격", "4. 멤버십 안내 (FREE / BASIC / PREMIUM)",
    "5. 플랫폼 이용수수료 (15%)", "6. 수업 진행 절차", "7. 기관 정보 제공", "8. 수강료 기준",
    "9. 플랫폼 이용 원칙", "10. 이용 제한", "11. HarmonyLink 지원 서비스", "12. 서비스 지역", "13. 문의"
]
for item in toc: add_bullet(item)

doc.add_page_break()
doc.add_heading("1. 운영 목적", level=1)
doc.add_paragraph("HarmonyLink는 교육기관과 전문 강사를 연결하고 수업 매칭 및 운영을 지원함으로써 공정하고 신뢰할 수 있는 교육서비스를 제공하는 것을 목표로 합니다.")
add_bullet("검토된 교육 파트너와 기관·수강자의 연결")
add_bullet("투명한 수업 조건과 수강료 기준 확정")
add_bullet("공정한 매칭 절차와 지속 가능한 파트너 생태계 운영")

doc.add_heading("2. 플랫폼 운영 방식", level=1)
doc.add_paragraph("HarmonyLink는 교육 매칭 플랫폼입니다. 플랫폼은 교육기관·수강자와 강사·교육업체를 연결하며, 수업료를 직접 수령하거나 보관하지 않습니다.")
add_bullet("모든 교육서비스는 입점 파트너가 직접 제공합니다.")
add_bullet("수업료는 기관 또는 수강자가 강사·교육업체에 직접 지급합니다.")
add_bullet("HarmonyLink는 프로그램 정보 확인, 매칭 승인, 일정 조율과 운영 지원을 제공합니다.")

doc.add_heading("3. 입점 파트너 자격", level=1)
doc.add_paragraph("입점 파트너는 관련 분야의 경력 또는 자격을 갖추고 책임감 있게 교육서비스를 제공할 수 있는 개인 강사 또는 교육업체입니다.")
add_bullet("HarmonyLink와 입점 파트너는 고용관계가 아닌 플랫폼 이용계약에 따른 독립적인 협력관계입니다.")
add_bullet("입점 파트너는 사업 운영, 보험, 세금 신고 및 관련 법적 의무를 스스로 부담합니다.")
add_bullet("지원서, 강사 프로필, 자격·경력 자료, 수업계획서 등을 제출해야 하며 프로그램에 따라 추가 자료가 필요할 수 있습니다.")

doc.add_page_break()
doc.add_heading("4. 입점 파트너 멤버십 안내", level=1)
doc.add_paragraph("HarmonyLink는 강사와 교육업체의 성장과 홍보를 지원하기 위해 FREE, BASIC, PREMIUM의 3단계 멤버십을 운영합니다. 멤버십은 입점 파트너를 위한 서비스이며 수강자의 플랫폼 가입·이용에는 회비가 없습니다.")

rows = [
    ("월 이용료", "무료", "$20 / 월", "$50 / 월"),
    ("기본 프로필·검색 노출", "●", "●", "●"),
    ("소형 배너 노출", "●", "●", "●"),
    ("기관·수강 의뢰 매칭", "가능 시", "기본", "우선"),
    ("교육 카테고리 노출", "1개", "전체", "전체 + 추천"),
    ("프로그램 수정", "분기 1회", "월 1회", "제한 없음"),
    ("운영 상담", "이메일", "기본 상담", "우선 상담"),
    ("입점 심사", "일반", "일반", "우선 심사"),
    ("메인·프리미엄 배너", "-", "-", "●"),
    ("하이벨 디자인 지원", "-", "-", "●"),
    ("스토어 수수료 할인", "-", "-", "●"),
    ("신규 기관 우선 매칭", "-", "-", "●"),
]
t = doc.add_table(rows=1, cols=4)
t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
widths=[Inches(2.65), Inches(1.25), Inches(1.25), Inches(1.55)]
headers=["혜택", "FREE", "BASIC", "PREMIUM"]
for i,(c,w,h) in enumerate(zip(t.rows[0].cells,widths,headers)):
    c.width=w; shade(c,NAVY); set_cell_text(c,h,True,WHITE,9,WD_ALIGN_PARAGRAPH.CENTER)
trPr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement("w:tblHeader"); rep.set(qn("w:val"),"true"); trPr.append(rep)
for ridx,row in enumerate(rows):
    cells=t.add_row().cells
    for i,(c,w,val) in enumerate(zip(cells,widths,row)):
        c.width=w; shade(c, PALE if ridx%2==0 else WHITE)
        set_cell_text(c,val,i==0,NAVY if i==0 else INK,8.2,WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("FREE 회원", level=2)
doc.add_paragraph("기본 등록으로 플랫폼을 체험하는 입점 파트너를 위한 무료 단계입니다.")
for x in ["업체·강사 기본 프로필 및 기본 프로그램 등록", "기본 검색과 소형 배너 노출", "플랫폼 공지 및 뉴스레터", "기관·수강 의뢰는 가능한 경우 안내"]: add_bullet(x)
doc.add_heading("BASIC 회원 ($20 / 월)", level=2)
doc.add_paragraph("기본적인 홍보와 정기적인 기관 매칭 안내가 필요한 파트너를 위한 멤버십입니다.")
for x in ["FREE 회원의 모든 혜택", "전체 관련 교육 카테고리 노출", "기관 및 수강 의뢰 매칭 안내", "프로그램 정보 월 1회 수정", "기본 심사 및 운영 상담"]: add_bullet(x)
doc.add_heading("PREMIUM 회원 ($50 / 월)", level=2)
doc.add_paragraph("브랜드 홍보와 기관 매칭을 적극적으로 확대하려는 파트너를 위한 프리미엄 멤버십입니다.")
for x in ["BASIC 회원의 모든 혜택", "하이벨 전문 디자이너의 홍보 전단·배너 디자인 지원", "전문교육 섹션과 메인·카테고리 추천 영역 우선 노출", "신규 기관 의뢰 우선 매칭 및 입점·프로그램 심사 우선 처리", "디지털 스토어 판매 수수료 할인", "운영 컨설팅과 특별 프로모션 우선 참여"]: add_bullet(x)
add_callout("멤버십과 수수료", "모든 멤버십은 운영 정책을 준수하는 입점 파트너에게 제공됩니다. 플랫폼 이용수수료 15%는 멤버십 종류와 관계없이 동일하게 적용됩니다.", fill="FFF5DD", color="8A5A00")

doc.add_page_break()
doc.add_heading("5. 플랫폼 이용수수료", level=1)
doc.add_paragraph("HarmonyLink는 플랫폼을 통해 연결된 모든 수업에 대해 HarmonyLink에 등록 및 확정된 수강료를 기준으로 15%의 플랫폼 이용수수료를 부과합니다.")
add_callout("납부 시점", "플랫폼 이용수수료는 수업 매칭 확정 후, 수업 진행 전에 선납해야 합니다. 납부가 완료되어야 최종 매칭이 승인되고 기관 상세정보가 제공됩니다.")
doc.add_heading("이용수수료 산정 기준", level=2)
doc.add_paragraph("수수료는 실제 할인 후 지급받는 금액이 아니라 HarmonyLink에 등록 및 확정된 수강료를 기준으로 계산합니다. 다음 할인에도 산정 기준은 변경되지 않습니다.")
for x in ["강사가 제공하는 할인", "이벤트·프로모션 할인", "기관 특별 할인", "쿠폰 및 기타 할인 혜택"]: add_bullet(x)

doc.add_heading("6. 수업 진행 절차", level=1)
steps = [
    ("STEP 1", "기관 또는 수강자가 HarmonyLink에 수업을 신청합니다."),
    ("STEP 2", "HarmonyLink가 프로그램과 조건에 적합한 입점 파트너를 연결합니다."),
    ("STEP 3", "프로그램, 등록 수강료, 횟수, 일정, 지역을 최종 확정합니다."),
    ("STEP 4", "입점 파트너가 플랫폼 이용수수료 15%를 선납합니다."),
    ("STEP 5", "입금 확인 후 기관 담당자 정보, 장소, 일정과 수업 승인번호를 제공합니다."),
    ("STEP 6", "입점 파트너와 기관·수강자가 직접 수업을 진행합니다."),
    ("STEP 7", "기관 또는 수강자가 입점 파트너에게 수업료를 직접 지급합니다."),
]
for label,text in steps:
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    t.columns[0].width=Inches(1.05); t.columns[1].width=Inches(5.65)
    c1,c2=t.rows[0].cells; shade(c1,BLUE); shade(c2,LIGHT)
    set_cell_text(c1,label,True,WHITE,8.5,WD_ALIGN_PARAGRAPH.CENTER); set_cell_text(c2,text,False,INK,9.2)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)

doc.add_page_break()
doc.add_heading("7. 기관 정보 제공", level=1)
doc.add_paragraph("기관 담당자 연락처, 이메일, 상세 주소 및 기타 매칭 정보는 플랫폼 이용수수료 납부 후 입점 파트너에게 제공합니다. 이는 기관 개인정보와 정상적인 매칭 절차를 보호하기 위한 정책입니다.")

doc.add_heading("8. 수강료 기준", level=1)
doc.add_paragraph("HarmonyLink를 통해 연결되는 모든 수업은 플랫폼에 등록 및 확정된 수강료를 기준으로 운영됩니다. 수업료 변경이 필요한 경우 수업 확정 전에 HarmonyLink와 협의하여 승인을 받아야 합니다.")
add_callout("결제 관계", "HarmonyLink는 수업료를 대신 수령하거나 보관하지 않습니다. 기관 또는 수강자가 입점 파트너에게 직접 지급합니다.")

doc.add_heading("9. 플랫폼 이용 원칙", level=1)
for x in [
    "HarmonyLink를 통해 소개된 기관과의 수업은 플랫폼을 통한 매칭으로 간주되며 이용수수료가 적용됩니다.",
    "입점 파트너는 정확한 프로필, 프로그램, 경력 및 수업 정보를 제공해야 합니다.",
    "확정된 일정과 수업 조건을 성실하게 이행하고 기관·수강자를 존중해야 합니다.",
    "기관 연락처와 매칭 정보는 해당 수업 운영 목적에만 사용해야 합니다.",
    "분쟁 또는 일정 변경이 발생하면 즉시 HarmonyLink에 알려야 합니다.",
]: add_bullet(x)

doc.add_heading("10. 이용 제한", level=1)
doc.add_paragraph("다음의 경우 플랫폼 이용이 제한되거나 파트너 자격이 일시 중지 또는 해지될 수 있습니다.")
for x in ["플랫폼 이용수수료 미납", "허위 정보·경력·자격 또는 수업 정보 제출", "플랫폼 정책 또는 이용계약 위반", "기관·수강자 또는 플랫폼의 신뢰를 훼손하는 행위", "기관 정보를 이용한 플랫폼 우회 거래", "관계 법령 위반 또는 안전상 중대한 문제"]: add_bullet(x)

doc.add_heading("11. HarmonyLink 지원 서비스", level=1)
for x in ["교육기관 및 수강 의뢰 매칭", "프로그램과 강사 홍보", "일정 조율과 운영 상담", "플랫폼 운영 지원", "신규 기관 발굴", "멤버십에 따른 디자인·배너·우선 노출 지원"]: add_bullet(x)

doc.add_heading("12. 서비스 지역", level=1)
for x in ["New York City", "Long Island", "Westchester", "New Jersey"]: add_bullet(x)
doc.add_paragraph("※ 향후 미국 동부 지역 및 온라인 서비스로 확대할 예정입니다.")

doc.add_heading("13. 문의", level=1)
contact = doc.add_table(rows=5, cols=2); contact.alignment=WD_TABLE_ALIGNMENT.CENTER; contact.autofit=False
contact.columns[0].width=Inches(1.45); contact.columns[1].width=Inches(5.25)
for i,(label,value) in enumerate([
    ("플랫폼", "HarmonyLink / 이음문화센터"), ("운영사", "Hibelle Consulting Inc."),
    ("홈페이지", "www.hibelleharmony.com"), ("이메일", "hibelle@hibelleconsulting.com"), ("대표전화", "929-603-0052")
]):
    c1,c2=contact.rows[i].cells; shade(c1,PALE); shade(c2,WHITE)
    set_cell_text(c1,label,True,NAVY,9); set_cell_text(c2,value,False,INK,9.5)

doc.add_page_break()
principle = doc.add_table(rows=1, cols=1)
principle.alignment = WD_TABLE_ALIGNMENT.CENTER
principle.autofit = False
principle.columns[0].width = Inches(6.7)
principle_cell = principle.cell(0, 0)
shade(principle_cell, NAVY)
cell_margins(principle_cell, 420, 420, 420, 420)

principle_title = principle_cell.paragraphs[0]
principle_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
principle_title.paragraph_format.space_after = Pt(18)
principle_title_run = principle_title.add_run("운영 원칙")
principle_title_run.bold = True
principle_title_run.font.size = Pt(22)
principle_title_run.font.color.rgb = RGBColor.from_string(WHITE)

for principle_text in [
    "Harmony Link는 강사와 교육기관이 함께 성장하는 신뢰 기반의 교육 플랫폼을 지향합니다.",
    "공정한 매칭, 투명한 운영, 상호 존중을 바탕으로 지속 가능한 교육 생태계를 만들어 갑니다."
]:
    principle_paragraph = principle_cell.add_paragraph()
    principle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    principle_paragraph.paragraph_format.space_after = Pt(13)
    principle_paragraph.paragraph_format.line_spacing = 1.35
    principle_run = principle_paragraph.add_run(principle_text)
    principle_run.bold = True
    principle_run.font.size = Pt(14)
    principle_run.font.color.rgb = RGBColor.from_string(WHITE)

doc.core_properties.title = "HarmonyLink 입점 파트너 플랫폼 이용 및 운영 정책"
doc.core_properties.subject = "Partner Platform Policy Version 1.0"
doc.core_properties.author = "Hibelle Consulting Inc."
doc.core_properties.keywords = "HarmonyLink, partner, policy, membership, platform fee"
doc.save(DOCX)
print(DOCX)
