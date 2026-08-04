from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "downloads"
OUT.mkdir(exist_ok=True)

NAVY = RGBColor(7, 38, 93)
BLUE = RGBColor(23, 104, 213)
MUTED = RGBColor(92, 108, 131)
LIGHT = "EAF5FF"
GOLD = "FFF4D6"
INK = RGBColor(23, 35, 58)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for side, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
        tc_mar.append(node)

def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = color

def setup(doc, label):
    section = doc.sections[0]
    section.top_margin = Inches(.72); section.bottom_margin = Inches(.72)
    section.left_margin = Inches(.82); section.right_margin = Inches(.82)
    section.header_distance = Inches(.35); section.footer_distance = Inches(.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.2
    for name, size in (("Heading 1",17),("Heading 2",13),("Heading 3",11.5)):
        style = doc.styles[name]; style.font.name = "Arial"; style.font.size = Pt(size)
        style.font.bold = True; style.font.color.rgb = NAVY
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.space_before = Pt(12); style.paragraph_format.space_after = Pt(6)
    hp = section.header.paragraphs[0]
    hp.text = f"HARMONYLINK  |  {label}"; hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in hp.runs: set_font(r, 8.5, True, BLUE)
    fp = section.footer.paragraphs[0]
    fp.text = "Hibelle Consulting Inc.  ·  Partner Center"; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs: set_font(r, 8, False, MUTED)

def title(doc, kicker, heading, subtitle):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    set_font(p.add_run(kicker.upper()), 9, True, BLUE)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    set_font(p.add_run(heading), 25, True, NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18)
    set_font(p.add_run(subtitle), 11, False, MUTED)

def callout(doc, label, text, fill=LIGHT):
    table=doc.add_table(rows=1, cols=1); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.autofit=False; table.columns[0].width=Inches(6.7)
    cell=table.cell(0,0); shade(cell,fill); margins(cell,130,160,130,160)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(3)
    set_font(p.add_run(label+"\n"),10,True,NAVY)
    set_font(p.add_run(text),9.5,False,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def bullet(doc, text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(4)
    set_font(p.add_run(text),10.5)

def numbered(doc, text):
    p=doc.add_paragraph(style="List Number"); p.paragraph_format.space_after=Pt(5)
    set_font(p.add_run(text),10.5)

def add_signature(doc):
    doc.add_heading("서명", level=1)
    table=doc.add_table(rows=5, cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.style="Table Grid"; table.autofit=False
    rows=[("회사","Hibelle Consulting Inc."),("대표자","____________________________"),("입점 파트너","____________________________"),("서명","____________________________"),("날짜","________년 ____월 ____일")]
    for row,(a,b) in zip(table.rows,rows):
        row.cells[0].width=Inches(1.45); row.cells[1].width=Inches(5.25)
        shade(row.cells[0],LIGHT); margins(row.cells[0]); margins(row.cells[1])
        row.cells[0].vertical_alignment=row.cells[1].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text=a; row.cells[1].text=b
        for r in row.cells[0].paragraphs[0].runs:set_font(r,9.5,True,NAVY)
        for r in row.cells[1].paragraphs[0].runs:set_font(r,9.5)

def build_contract():
    doc=Document(); setup(doc,"PARTNER AGREEMENT"); title(doc,"Partner Agreement","HarmonyLink 입점 파트너 계약서","운영용 기본 계약서 초안 · Version 1.0")
    callout(doc,"중요 안내","본 문서는 운영을 위한 기본 계약서 초안입니다. 실제 서명 전 당사자 정보, 서비스 범위, 관할법과 세부 조건을 확인하고 필요하면 변호사의 검토를 받으세요.",GOLD)
    doc.add_heading("1. 계약 당사자",level=1)
    doc.add_paragraph("본 계약은 Hibelle Consulting Inc.(이하 ‘회사’)와 아래 서명한 교육업체 또는 개인 강사(이하 ‘파트너’) 사이의 HarmonyLink 플랫폼 입점 및 프로그램 운영에 관한 기본 사항을 정합니다.")
    doc.add_heading("2. 계약 목적과 범위",level=1)
    bullet(doc,"회사는 파트너의 프로필과 승인된 프로그램을 HarmonyLink에 게시하고 문의·매칭을 지원합니다.")
    bullet(doc,"파트너는 정확한 정보와 적법하게 사용할 수 있는 이미지·교육자료를 제공하고 합의된 수업을 성실히 진행합니다.")
    bullet(doc,"구체적인 일정, 장소, 비용과 준비사항은 각 수업 또는 기관과의 별도 협의로 확정합니다.")
    doc.add_heading("3. 멤버십",level=1)
    table=doc.add_table(rows=3,cols=3); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    data=[("구분","월 등록비","기본 범위"),("BASIC 승인 파트너","$20","프로필·프로그램 등록, 기본 자료실, 문의 연결"),("PREMIUM 승인 파트너","$50","BASIC 혜택과 PREMIUM 자료·홍보 신청 기능")]
    widths=[Inches(1.75),Inches(1.15),Inches(3.8)]
    for ri,row in enumerate(table.rows):
        for ci,cell in enumerate(row.cells):
            cell.width=widths[ci]; margins(cell); cell.text=data[ri][ci]
            if ri==0: shade(cell,"DCEBFF")
            for r in cell.paragraphs[0].runs:set_font(r,9.2,ri==0,NAVY if ri==0 else INK)
    doc.add_paragraph("멤버십 적용일, 결제 방식, 변경 또는 해지 시점은 별도 안내 또는 서면 합의에 따릅니다.")
    doc.add_heading("4. 수수료와 정산",level=1)
    bullet(doc,"플랫폼을 통해 성사된 유료 수업은 현재 운영정책에 따른 플랫폼 이용수수료가 적용될 수 있습니다.")
    bullet(doc,"구체적인 수수료율, 결제 주체, 환불과 정산일은 수업 확정 전에 서면 또는 전자메시지로 확인합니다.")
    doc.add_heading("5. 파트너의 의무",level=1)
    for text in ["자격·경력·가격·프로그램 내용을 사실대로 제공합니다.","고객 개인정보를 수업 진행 목적에만 사용하고 외부에 공유하지 않습니다.","안전수칙, 기관 규정, 저작권과 관련 법규를 준수합니다.","일정 변경, 사고, 민원 또는 수업 진행 문제가 발생하면 회사에 신속히 알립니다."] : bullet(doc,text)
    doc.add_heading("6. 취소·변경·환불",level=1)
    doc.add_paragraph("수업별 취소·변경·환불 조건은 예약 확정 전에 고객, 파트너와 회사가 확인한 조건을 우선 적용합니다. 별도 합의가 없으면 회사의 최신 운영정책을 따릅니다.")
    doc.add_heading("7. 지식재산권과 홍보자료",level=1)
    doc.add_paragraph("각 당사자가 기존에 보유한 상표, 로고, 교안과 콘텐츠의 권리는 원 소유자에게 있습니다. 파트너는 입점 기간 동안 회사가 승인된 프로필과 홍보자료를 HarmonyLink 홍보 목적으로 게시할 수 있도록 허용합니다.")
    doc.add_heading("8. 계약기간과 종료",level=1)
    doc.add_paragraph("계약은 서명일에 시작하며 어느 당사자든 서면 또는 이메일로 종료를 요청할 수 있습니다. 미정산 금액, 개인정보 보호, 저작권과 비밀유지 의무는 종료 후에도 필요한 범위에서 유지됩니다.")
    doc.add_heading("9. 책임과 분쟁",level=1)
    doc.add_paragraph("당사자는 문제 발생 시 우선 성실히 협의합니다. 고의 또는 중대한 과실, 허위 정보, 권리 침해 등 각 당사자의 귀책 사유로 발생한 손해는 해당 당사자가 책임집니다. 관할과 준거법은 최종 서명 전 당사자가 별도로 확인합니다.")
    doc.add_heading("10. 전체 합의",level=1)
    doc.add_paragraph("본 계약과 참조된 최신 운영정책은 입점 운영에 관한 기본 합의를 구성합니다. 중요한 변경은 기록이 남는 방식으로 합의합니다.")
    add_signature(doc)
    doc.core_properties.title="HarmonyLink 입점 파트너 계약서"; doc.core_properties.author="Hibelle Consulting Inc."
    doc.save(OUT/"HarmonyLink_Partner_Agreement_Draft_v1.0.docx")

def build_guide():
    doc=Document(); setup(doc,"INSTRUCTOR GUIDE"); title(doc,"Partner Operations","HarmonyLink 강사 활동 가이드","첫 등록부터 수업 완료까지 필요한 실무 기준")
    callout(doc,"핵심 원칙","정확한 안내, 빠른 응답, 안전한 진행, 개인정보 보호를 모든 수업의 기본으로 합니다.")
    for heading,items in [
        ("1. 활동 시작 전",["프로필, 연락처, 활동 지역과 가능한 시간을 최신 상태로 유지합니다.","수업명, 대상, 시간, 비용, 준비물과 최소·최대 인원을 명확하게 정리합니다.","강사 사진과 전단지는 본인이 사용할 권리가 있는 자료만 제출합니다."]),
        ("2. 문의를 받았을 때",["가능하면 1영업일 안에 응답합니다.","날짜, 장소, 인원, 대상 연령, 예산, 필요한 장비를 확인합니다.","확정되지 않은 일정이나 효과를 보장하지 않습니다."]),
        ("3. 수업 확정 후",["고객과 최종 일정·주소·주차·담당자 연락처를 다시 확인합니다.","교안, 장비, 출석부와 비상 연락 수단을 준비합니다.","변경이 필요하면 가능한 한 빨리 고객과 HarmonyLink에 알립니다."]),
        ("4. 수업 당일",["약속 시간보다 여유 있게 도착하고 기관의 출입·안전 규정을 따릅니다.","참여자를 존중하며 차별적 표현, 강압적 판매와 허위 광고를 하지 않습니다.","사진과 영상은 사전 동의를 받은 경우에만 촬영·사용합니다."]),
        ("5. 수업 종료 후",["출석과 진행 결과, 특이사항을 정리합니다.","정산에 필요한 정보를 정해진 방식으로 제출합니다.","개인정보가 포함된 자료는 안전하게 보관하고 불필요해지면 삭제합니다."]),
        ("6. 문제 발생 시",["안전이 우선이며 필요한 경우 수업을 중단하고 기관 담당자에게 알립니다.","민원, 사고, 일정 충돌 또는 결제 문제가 생기면 임의로 확정하지 말고 HarmonyLink와 협의합니다."])
    ]:
        doc.add_heading(heading,level=1)
        for item in items: bullet(doc,item)
    doc.add_heading("빠른 체크리스트",level=1)
    for item in ["□ 일정과 장소 확인","□ 담당자 연락처 확인","□ 교안·장비·준비물 확인","□ 안전 및 촬영 동의 확인","□ 수업 후 결과와 정산 자료 제출"]: doc.add_paragraph(item)
    doc.core_properties.title="HarmonyLink 강사 활동 가이드"; doc.core_properties.author="Hibelle Consulting Inc."
    doc.save(OUT/"HarmonyLink_Instructor_Activity_Guide_v1.0.docx")

def build_faq():
    doc=Document(); setup(doc,"PARTNER FAQ"); title(doc,"Partner Help","HarmonyLink 입점 파트너 FAQ","자주 묻는 질문과 빠른 답변")
    faqs=[
      ("누가 입점 파트너가 될 수 있나요?","교육 프로그램을 보유한 업체, 단체, 개인 강사가 신청할 수 있으며 검토와 승인 후 활동합니다."),
      ("$20 파트너와 $50 파트너의 차이는 무엇인가요?","$20 BASIC 파트너는 기본 등록·자료실·문의 연결 기능을 이용합니다. $50 PREMIUM 파트너는 기본 혜택에 더해 PREMIUM 자료와 추가 홍보 신청 기능을 이용할 수 있습니다."),
      ("프로그램은 어떻게 등록하나요?","수업명, 대상, 시간, 비용, 지역, 설명, 대표 이미지와 강사 정보를 준비해 HarmonyLink에 전달합니다. 검토 후 게시됩니다."),
      ("가격과 일정은 누가 정하나요?","파트너가 기본 조건을 제시하고 고객·기관의 요청을 확인한 뒤 최종 조건을 협의합니다."),
      ("문의가 오면 얼마나 빨리 답해야 하나요?","가능하면 1영업일 안에 응답해 주세요. 답변이 늦어질 때는 예상 응답 시간을 먼저 알려 주세요."),
      ("수업을 취소하거나 변경해야 하면 어떻게 하나요?","고객과 HarmonyLink에 즉시 알리고, 예약 때 확인한 취소·변경 조건에 따라 대체 일정 또는 환불을 협의합니다."),
      ("고객 연락처를 개인 홍보에 사용해도 되나요?","안 됩니다. 고객 개인정보는 해당 문의와 수업 진행 목적에만 사용해야 하며 별도 동의 없이 홍보나 외부 공유에 사용할 수 없습니다."),
      ("사진과 영상을 홍보에 사용할 수 있나요?","참여자와 기관의 사전 동의를 받은 자료만 사용할 수 있습니다. 아동·청소년 자료는 보호자 및 기관 규정을 반드시 확인하세요."),
      ("정산은 어떻게 진행되나요?","수업별 결제 방식, 수수료와 정산일을 예약 확정 전에 안내합니다. 필요한 출석·완료 자료를 제출해야 할 수 있습니다."),
      ("멤버십을 변경하거나 종료하려면 어떻게 하나요?","문의하기를 통해 요청해 주세요. 적용일과 미정산 내역을 확인한 뒤 변경 또는 종료를 처리합니다."),
      ("도움이 필요하면 어디로 연락하나요?","hibelle@hibelleconsulting.com 또는 미국 상담 929-603-0052로 연락해 주세요.")
    ]
    for i,(q,a) in enumerate(faqs,1):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(9); p.paragraph_format.space_after=Pt(3)
        set_font(p.add_run(f"Q{i}. {q}"),11.5,True,NAVY)
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.space_after=Pt(8)
        set_font(p.add_run(a),10.5)
    callout(doc,"문의","이 문서에서 해결되지 않는 내용은 홈페이지 문의하기 또는 hibelle@hibelleconsulting.com으로 보내 주세요.")
    doc.core_properties.title="HarmonyLink 입점 파트너 FAQ"; doc.core_properties.author="Hibelle Consulting Inc."
    doc.save(OUT/"HarmonyLink_Partner_FAQ_v1.0.docx")

if __name__ == "__main__":
    build_contract(); build_guide(); build_faq()
